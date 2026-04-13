"""Document processor for various file types.

This module handles loading and processing different document formats
for ingestion into the vector database.
"""

from __future__ import annotations

import os
import hashlib
from abc import ABC, abstractmethod
from pathlib import Path
from typing import BinaryIO, List, Optional, Dict, Any
from dataclasses import dataclass
import io

from langchain_core.documents import Document


@dataclass
class ProcessedDocument:
    """Represents a processed document with metadata."""
    
    content: str
    metadata: Dict[str, Any]
    source: str
    doc_type: str


class BaseDocumentLoader(ABC):
    """Base class for document loaders."""
    
    supported_extensions: List[str] = []
    
    @abstractmethod
    def load(self, file_content: bytes, filename: str, **kwargs) -> List[ProcessedDocument]:
        """Load documents from file content.
        
        Args:
            file_content: Raw file content as bytes.
            filename: Original filename.
            **kwargs: Additional loader-specific arguments.
            
        Returns:
            List of processed documents.
        """
        pass
    
    def can_load(self, filename: str) -> bool:
        """Check if this loader can handle the given filename.
        
        Args:
            filename: Name of the file.
            
        Returns:
            True if this loader supports the file type.
        """
        ext = Path(filename).suffix.lower()
        return ext in self.supported_extensions


class TextDocumentLoader(BaseDocumentLoader):
    """Loader for plain text files."""
    
    supported_extensions = [".txt", ".md", ".markdown", ".json", ".csv"]
    
    def load(self, file_content: bytes, filename: str, **kwargs) -> List[ProcessedDocument]:
        """Load text document."""
        # Try different encodings
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
        text = None
        
        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        
        if text is None:
            raise ValueError(f"Could not decode file {filename} with any supported encoding")
        
        # Split into chunks if content is large
        chunk_size = kwargs.get("chunk_size", 2000)
        overlap = kwargs.get("overlap", 200)
        
        if len(text) <= chunk_size:
            return [ProcessedDocument(
                content=text,
                metadata={
                    "source": filename,
                    "chunk_index": 0,
                    "total_chunks": 1,
                },
                source=filename,
                doc_type="text"
            )]
        
        # Chunk the text
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = start + chunk_size
            chunk_text = text[start:end]
            
            chunks.append(ProcessedDocument(
                content=chunk_text,
                metadata={
                    "source": filename,
                    "chunk_index": chunk_index,
                    "total_chunks": None,  # Will update later
                    "char_start": start,
                    "char_end": end,
                },
                source=filename,
                doc_type="text"
            ))
            
            start = end - overlap
            chunk_index += 1
        
        # Update total chunks
        for chunk in chunks:
            chunk.metadata["total_chunks"] = len(chunks)
        
        return chunks


class PDFDocumentLoader(BaseDocumentLoader):
    """Loader for PDF files."""
    
    supported_extensions = [".pdf"]
    
    def load(self, file_content: bytes, filename: str, **kwargs) -> List[ProcessedDocument]:
        """Load PDF document."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError(
                "pypdf is required for PDF processing. "
                "Install it with: pip install pypdf"
            )
        
        reader = PdfReader(io.BytesIO(file_content))
        documents = []
        
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text.strip():
                documents.append(ProcessedDocument(
                    content=text,
                    metadata={
                        "source": filename,
                        "page_number": page_num,
                        "total_pages": len(reader.pages),
                    },
                    source=filename,
                    doc_type="pdf"
                ))
        
        return documents


class WordDocumentLoader(BaseDocumentLoader):
    """Loader for Word documents."""
    
    supported_extensions = [".docx"]  # .doc 需要额外处理
    
    def load(self, file_content: bytes, filename: str, **kwargs) -> List[ProcessedDocument]:
        """Load Word document."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError(
                "python-docx is required for Word processing. "
                "Install it with: pip install python-docx"
            )
        
        # 检查文件头 - docx 是 zip 格式
        if not file_content.startswith(b'PK'):
            raise ValueError(
                f"文件 {filename} 不是有效的 .docx 格式。"
                f"如果是 .doc 旧格式，请先用 Word 另存为 .docx"
            )
        
        try:
            doc = DocxDocument(io.BytesIO(file_content))
        except Exception as e:
            if "not a zip file" in str(e).lower():
                raise ValueError(
                    f"文件 {filename} 损坏或不是有效的 .docx 格式。"
                    f"提示: .doc 旧格式需要转换为 .docx"
                )
            raise
        
        # Extract text from all paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Join paragraphs and create chunks
        full_text = "\n\n".join(paragraphs)
        
        chunk_size = kwargs.get("chunk_size", 2000)
        overlap = kwargs.get("overlap", 200)
        
        if len(full_text) <= chunk_size:
            return [ProcessedDocument(
                content=full_text,
                metadata={
                    "source": filename,
                    "chunk_index": 0,
                    "total_chunks": 1,
                    "paragraph_count": len(paragraphs),
                },
                source=filename,
                doc_type="word"
            )]
        
        # Create chunks
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(full_text):
            end = min(start + chunk_size, len(full_text))
            # Try to break at paragraph boundary
            if end < len(full_text):
                next_para = full_text.find("\n\n", end - 100, end + 100)
                if next_para != -1:
                    end = next_para + 2
            
            chunk_text = full_text[start:end]
            chunks.append(ProcessedDocument(
                content=chunk_text,
                metadata={
                    "source": filename,
                    "chunk_index": chunk_index,
                    "total_chunks": None,
                    "char_start": start,
                },
                source=filename,
                doc_type="word"
            ))
            
            start = end - overlap
            chunk_index += 1
        
        # Update total chunks
        for chunk in chunks:
            chunk.metadata["total_chunks"] = len(chunks)
            chunk.metadata["paragraph_count"] = len(paragraphs)
        
        return chunks


class HTMLDocumentLoader(BaseDocumentLoader):
    """Loader for HTML files."""
    
    supported_extensions = [".html", ".htm"]
    
    def load(self, file_content: bytes, filename: str, **kwargs) -> List[ProcessedDocument]:
        """Load HTML document."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError(
                "beautifulsoup4 is required for HTML processing. "
                "Install it with: pip install beautifulsoup4"
            )
        
        # Decode content
        encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
        text = None
        
        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        
        if text is None:
            raise ValueError(f"Could not decode HTML file {filename}")
        
        soup = BeautifulSoup(text, "html.parser")
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text content
        text = soup.get_text(separator="\n", strip=True)
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = "\n".join(chunk for chunk in chunks if chunk)
        
        return [ProcessedDocument(
            content=text,
            metadata={
                "source": filename,
                "title": soup.title.string if soup.title else filename,
            },
            source=filename,
            doc_type="html"
        )]


class DocumentProcessor:
    """Main document processor that manages multiple loaders."""
    
    def __init__(self):
        """Initialize document processor with default loaders."""
        self.loaders: List[BaseDocumentLoader] = [
            TextDocumentLoader(),
            PDFDocumentLoader(),
            WordDocumentLoader(),
            HTMLDocumentLoader(),
        ]
    
    def register_loader(self, loader: BaseDocumentLoader) -> None:
        """Register a new document loader.
        
        Args:
            loader: Document loader instance.
        """
        self.loaders.append(loader)
    
    def process_file(
        self, 
        file_content: bytes, 
        filename: str,
        **kwargs
    ) -> List[ProcessedDocument]:
        """Process a file and return documents.
        
        Args:
            file_content: Raw file content.
            filename: Original filename.
            **kwargs: Additional processing arguments.
            
        Returns:
            List of processed documents.
            
        Raises:
            ValueError: If no suitable loader is found.
        """
        for loader in self.loaders:
            if loader.can_load(filename):
                return loader.load(file_content, filename, **kwargs)
        
        raise ValueError(f"No loader found for file: {filename}")
    
    def process_files(
        self,
        files: List[Dict[str, Any]],
        **kwargs
    ) -> List[ProcessedDocument]:
        """Process multiple files.
        
        Args:
            files: List of file dictionaries with 'content' and 'filename' keys.
            **kwargs: Additional processing arguments.
            
        Returns:
            List of processed documents from all files.
        """
        all_documents = []
        
        for file_info in files:
            content = file_info.get("content")
            filename = file_info.get("filename")
            
            if not content or not filename:
                continue
            
            try:
                docs = self.process_file(content, filename, **kwargs)
                all_documents.extend(docs)
            except Exception as e:
                # Log error but continue processing other files
                print(f"Error processing {filename}: {e}")
                continue
        
        return all_documents
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions.
        
        Returns:
            List of supported extensions.
        """
        extensions = []
        for loader in self.loaders:
            extensions.extend(loader.supported_extensions)
        return extensions


# Global processor instance
_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """Get or create global document processor.
    
    Returns:
        DocumentProcessor instance.
    """
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor


def processed_docs_to_langchain(docs: List[ProcessedDocument]) -> List[Document]:
    """Convert processed documents to LangChain Document format.
    
    Args:
        docs: List of processed documents.
        
    Returns:
        List of LangChain Document objects.
    """
    return [
        Document(
            page_content=doc.content,
            metadata=doc.metadata
        )
        for doc in docs
    ]
